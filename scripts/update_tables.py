import sys
import os
from docx import Document
from docx.shared import Pt

def update_tables(docx_path):
    # Ensure the docx_path is an absolute path
    docx_path = os.path.abspath(docx_path)

    # Check if the file exists
    if not os.path.exists(docx_path):
        raise FileNotFoundError(f"Error: Cannot find DOCX file at '{docx_path}'")

    # Load the document
    doc = Document(docx_path)

    # Update the tables
    for table in doc.tables:
        table.autofit = True
        # Set smaller font size for all cells in the table
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        # Set font size to 9pt (slightly smaller than default 11-12pt)
                        run.font.size = Pt(9)

    # Save the updated document
    doc.save(docx_path)

if __name__ == "__main__":
    if len(sys.argv) != 2:
        print("Usage: python3 update_tables.py <docx_path>")
        sys.exit(1)

    update_tables(sys.argv[1])
